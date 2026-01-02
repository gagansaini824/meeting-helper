import os
import io
import base64
import json
import pickle
from typing import Optional
from dataclasses import dataclass, field, asdict
from datetime import datetime
import numpy as np
from openai import OpenAI
from pypdf import PdfReader
from docx import Document as DocxDocument
import openpyxl
from PIL import Image
import anthropic

@dataclass
class DocumentChunk:
    id: str
    document_id: str
    content: str
    embedding: Optional[list[float]] = None

@dataclass
class Document:
    id: str
    name: str
    content: str
    chunks: list[DocumentChunk] = field(default_factory=list)
    uploaded_at: datetime = field(default_factory=datetime.now)

class DocumentStore:
    def __init__(self, storage_dir: str = "data"):
        self.documents: dict[str, Document] = {}
        self._openai: Optional[OpenAI] = None
        self._anthropic: Optional[anthropic.Anthropic] = None
        self.storage_dir = storage_dir
        self.db_file = os.path.join(storage_dir, "documents.pkl")

        # Create storage directory if it doesn't exist
        os.makedirs(storage_dir, exist_ok=True)

        # Load existing documents from disk
        self._load_from_disk()

    def _get_openai(self) -> OpenAI:
        if self._openai is None:
            self._openai = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
        return self._openai

    def _get_anthropic(self) -> anthropic.Anthropic:
        if self._anthropic is None:
            self._anthropic = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
        return self._anthropic

    def _extract_text_from_file(self, file_bytes: bytes, filename: str) -> str:
        """Extract text from various file formats"""
        extension = filename.lower().split('.')[-1]

        try:
            # PDF files
            if extension == 'pdf':
                pdf_reader = PdfReader(io.BytesIO(file_bytes))
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text.strip()

            # DOCX files
            elif extension == 'docx':
                doc = DocxDocument(io.BytesIO(file_bytes))
                text = "\n".join([para.text for para in doc.paragraphs])
                return text.strip()

            # Markdown files
            elif extension == 'md':
                return file_bytes.decode('utf-8')

            # CSV files
            elif extension == 'csv':
                import csv
                text = file_bytes.decode('utf-8')
                reader = csv.reader(io.StringIO(text))
                rows = [', '.join(row) for row in reader]
                return '\n'.join(rows)

            # Excel files
            elif extension in ['xlsx', 'xls']:
                wb = openpyxl.load_workbook(io.BytesIO(file_bytes))
                text = ""
                for sheet_name in wb.sheetnames:
                    sheet = wb[sheet_name]
                    text += f"\n=== Sheet: {sheet_name} ===\n"
                    for row in sheet.iter_rows(values_only=True):
                        text += ', '.join([str(cell) if cell is not None else '' for cell in row]) + '\n'
                return text.strip()

            # Image files - use Claude vision to analyze
            elif extension in ['png', 'jpg', 'jpeg', 'gif', 'bmp', 'webp']:
                return self._analyze_image(file_bytes, extension)

            # Plain text files
            else:
                return file_bytes.decode('utf-8')

        except Exception as e:
            raise ValueError(f"Error parsing {extension} file: {str(e)}")

    def _analyze_image(self, image_bytes: bytes, extension: str) -> str:
        """Analyze image using Claude vision"""
        try:
            # Convert image to base64
            image_b64 = base64.b64encode(image_bytes).decode('utf-8')

            # Determine media type
            media_type_map = {
                'png': 'image/png',
                'jpg': 'image/jpeg',
                'jpeg': 'image/jpeg',
                'gif': 'image/gif',
                'webp': 'image/webp'
            }
            media_type = media_type_map.get(extension, 'image/jpeg')

            client = self._get_anthropic()
            response = client.messages.create(
                model="claude-sonnet-4-5-20250929",
                max_tokens=2000,
                messages=[{
                    "role": "user",
                    "content": [
                        {
                            "type": "image",
                            "source": {
                                "type": "base64",
                                "media_type": media_type,
                                "data": image_b64
                            }
                        },
                        {
                            "type": "text",
                            "text": "Analyze this image in detail. Describe what you see, any text present, diagrams, charts, or other relevant information. Provide a comprehensive description that can be used for semantic search."
                        }
                    ]
                }]
            )

            return response.content[0].text

        except Exception as e:
            return f"[Image file - analysis failed: {str(e)}]"
    
    def _chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> list[str]:
        chunks = []
        start = 0
        while start < len(text):
            end = min(start + chunk_size, len(text))
            chunks.append(text[start:end])
            start = end - overlap
            if start + overlap >= len(text):
                break
        return chunks
    
    async def add_document(self, name: str, content: str, file_bytes: Optional[bytes] = None) -> Document:
        import uuid

        # If file_bytes provided, extract text from the file
        if file_bytes:
            content = self._extract_text_from_file(file_bytes, name)

        doc_id = f"doc_{int(datetime.now().timestamp())}_{uuid.uuid4().hex[:8]}"
        text_chunks = self._chunk_text(content)
        
        chunks = []
        for i, chunk_text in enumerate(text_chunks):
            chunk_id = f"{doc_id}_chunk_{i}"
            embedding = None
            
            try:
                response = self._get_openai().embeddings.create(
                    model="text-embedding-3-small",
                    input=chunk_text
                )
                embedding = response.data[0].embedding
            except Exception as e:
                print(f"Embedding error: {e}")
            
            chunks.append(DocumentChunk(
                id=chunk_id,
                document_id=doc_id,
                content=chunk_text,
                embedding=embedding
            ))
        
        doc = Document(
            id=doc_id,
            name=name,
            content=content,
            chunks=chunks
        )
        self.documents[doc_id] = doc
        self._save_to_disk()
        return doc

    def remove_document(self, doc_id: str) -> bool:
        if doc_id in self.documents:
            del self.documents[doc_id]
            self._save_to_disk()
            return True
        return False

    def _save_to_disk(self):
        """Save documents to disk"""
        try:
            with open(self.db_file, 'wb') as f:
                pickle.dump(self.documents, f)
        except Exception as e:
            print(f"Error saving documents: {e}")

    def _load_from_disk(self):
        """Load documents from disk"""
        try:
            if os.path.exists(self.db_file):
                with open(self.db_file, 'rb') as f:
                    self.documents = pickle.load(f)
                print(f"Loaded {len(self.documents)} documents from disk")
        except Exception as e:
            print(f"Error loading documents: {e}")
            self.documents = {}
    
    def list_documents(self) -> list[dict]:
        return [
            {"id": d.id, "name": d.name, "uploaded_at": d.uploaded_at.isoformat()}
            for d in self.documents.values()
        ]
    
    def _cosine_similarity(self, a: list[float], b: list[float]) -> float:
        a_np = np.array(a)
        b_np = np.array(b)
        return float(np.dot(a_np, b_np) / (np.linalg.norm(a_np) * np.linalg.norm(b_np)))
    
    async def search(self, query: str, top_k: int = 5) -> list[dict]:
        if not self.documents:
            return []
        
        try:
            response = self._get_openai().embeddings.create(
                model="text-embedding-3-small",
                input=query
            )
            query_embedding = response.data[0].embedding
        except Exception:
            return self._keyword_search(query, top_k)
        
        results = []
        for doc in self.documents.values():
            for chunk in doc.chunks:
                if chunk.embedding:
                    score = self._cosine_similarity(query_embedding, chunk.embedding)
                    results.append({
                        "chunk": chunk,
                        "score": score,
                        "document_name": doc.name
                    })
        
        results.sort(key=lambda x: x["score"], reverse=True)
        return results[:top_k]
    
    def _keyword_search(self, query: str, top_k: int) -> list[dict]:
        query_words = query.lower().split()
        results = []
        
        for doc in self.documents.values():
            for chunk in doc.chunks:
                chunk_lower = chunk.content.lower()
                score = sum(1 for word in query_words if word in chunk_lower) / len(query_words)
                if score > 0:
                    results.append({
                        "chunk": chunk,
                        "score": score,
                        "document_name": doc.name
                    })
        
        results.sort(key=lambda x: x["score"], reverse=True)
        return results[:top_k]
    
    def get_context(self, search_results: list[dict]) -> str:
        if not search_results:
            return ""
        
        parts = [
            f"[From: {r['document_name']}]\n{r['chunk'].content}"
            for r in search_results
        ]
        return "\n\n---\n\n".join(parts)

# Singleton
document_store = DocumentStore()
